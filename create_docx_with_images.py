#!/usr/bin/env python3
"""
使用python-docx库将images目录中的图片插入到DOCX文档中的示例
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def add_page_break_before(paragraph):
    """在段落前添加分页符"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pPr.insert_element_before(br,
        'w:ind', 'w:spacing', 'w:jc', 'w:keepLines', 'w:keepNext', 'w:widowControl')

def create_document_with_images():
    """
    创建包含所有Mermaid图片的DOCX文档
    """
    # 设置路径
    images_dir = Path("D:/ProjectPackage/DocGenerator/images")
    output_file = "D:/ProjectPackage/DocGenerator/mermaid_document_with_images.docx"

    # 创建新的Word文档
    doc = Document()

    # 添加标题
    title = doc.add_heading('Mermaid图表集合文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加简介
    intro_para = doc.add_paragraph()
    intro_para.add_run('本文档包含了由Mermaid语法生成的所有图表。').bold = True
    intro_para.add_run(' 这些图表是通过Python脚本自动从images目录中获取并插入到此文档中的。')
    doc.add_paragraph()  # 空行

    # 获取所有PNG图片文件
    if not images_dir.exists():
        print(f"错误：图片目录不存在: {images_dir}")
        return

    image_files = sorted(images_dir.glob("*.png"))

    if not image_files:
        print("错误：在images目录中没有找到PNG图片文件")
        return

    print(f"找到 {len(image_files)} 个图片文件")

    # 为每个图片创建一个章节
    for i, image_path in enumerate(image_files, 1):
        print(f"正在处理图片 {i}: {image_path.name}")

        # 添加分页符（除了第一个图片）
        if i > 1:
            doc.add_page_break()

        # 章节标题
        section_title = doc.add_heading(f'图表 {i}: {image_path.stem}', level=1)
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图片信息
        info_para = doc.add_paragraph()
        info_para.add_run('图片文件名: ').bold = True
        info_para.add_run(image_path.name)

        # 获取图片文件大小
        file_size = image_path.stat().st_size
        info_para = doc.add_paragraph()
        info_para.add_run('文件大小: ').bold = True
        info_para.add_run(f'{file_size} 字节')

        doc.add_paragraph()  # 空行

        # 添加图片到文档
        try:
            # 设置图片宽度为6英寸，高度会自动按比例调整
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(6.0))

            # 添加图片说明
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.add_run(f'图 {i}: {image_path.stem.replace("-", " ").title()}').italic = True

            doc.add_paragraph()  # 空行

        except Exception as e:
            error_para = doc.add_paragraph()
            error_para.add_run(f'错误：无法插入图片 {image_path.name}: {str(e)}').italic = True

    # 添加结尾
    doc.add_page_break()
    conclusion = doc.add_heading('文档说明', level=1)
    doc.add_paragraph('此文档由Python脚本自动生成，使用了以下技术：')
    doc.add_paragraph('• Mermaid CLI (mmdc) - 用于生成图表图片', style='List Bullet')
    doc.add_paragraph('• python-docx - 用于创建和操作Word文档', style='List Bullet')
    doc.add_paragraph('• 项目图片目录: D:/ProjectPackage/DocGenerator/images/', style='List Bullet')

    # 保存文档
    try:
        doc.save(output_file)
        print(f"文档已成功保存到: {output_file}")

        # 显示文件信息
        output_path = Path(output_file)
        if output_path.exists():
            file_size = output_path.stat().st_size
            print(f"文档大小: {file_size} 字节")
            print(f"创建时间: {output_path.stat().st_ctime}")

        return True

    except Exception as e:
        print(f"保存文档时出错: {e}")
        return False

def main():
    """主函数"""
    print("=" * 60)
    print("开始创建包含Mermaid图片的DOCX文档")
    print("=" * 60)

    # 检查python-docx是否安装
    try:
        import docx
        print(f"python-docx 版本: {docx.__version__}")
    except ImportError:
        print("错误：未安装python-docx库")
        print("请运行: pip install python-docx")
        return

    # 创建文档
    success = create_document_with_images()

    if success:
        print("\n" + "=" * 60)
        print("文档创建完成！")
        print("=" * 60)
    else:
        print("\n" + "=" * 60)
        print("文档创建失败！")
        print("=" * 60)

if __name__ == "__main__":
    main()