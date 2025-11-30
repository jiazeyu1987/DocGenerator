#!/usr/bin/env python3
"""
简单示例：将项目images目录中的图片插入到DOCX文档中
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def simple_docx_example():
    """创建简单的图片文档示例"""

    # 设置路径
    images_dir = Path("D:/ProjectPackage/DocGenerator/images")
    output_file = "D:/ProjectPackage/DocGenerator/simple_mermaid_doc.docx"

    # 创建新文档
    doc = Document()

    # 添加标题
    title = doc.add_heading('Mermaid Flowcharts', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加说明
    doc.add_paragraph('This document demonstrates how to programmatically insert Mermaid diagram images into a Word document using python-docx.')
    doc.add_paragraph()

    # 获取所有PNG图片
    if not images_dir.exists():
        print(f"Images directory not found: {images_dir}")
        return False

    image_files = sorted(images_dir.glob("*.png"))

    if not image_files:
        print("No PNG images found in images directory")
        return False

    print(f"Found {len(image_files)} images to insert")

    # 为每个图片添加到文档
    for i, image_path in enumerate(image_files, 1):
        print(f"Inserting image {i}: {image_path.name}")

        # 添加图片标题
        heading = doc.add_heading(f'Flowchart {i}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图片（6英寸宽，自动调整高度）
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.0))

        # 添加图片说明
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(f'Figure {i}: {image_path.stem.replace("-", " ").title()}').italic = True

        # 添加空行
        doc.add_paragraph()

    # 保存文档
    try:
        doc.save(output_file)
        print(f"Document saved successfully: {output_file}")

        # 显示文件信息
        output_path = Path(output_file)
        if output_path.exists():
            size_kb = output_path.stat().st_size / 1024
            print(f"File size: {size_kb:.1f} KB")

        return True

    except Exception as e:
        print(f"Error saving document: {e}")
        return False

if __name__ == "__main__":
    print("Simple DOCX Example with Mermaid Images")
    print("=" * 50)

    success = simple_docx_example()

    if success:
        print("Success! Check the generated DOCX file.")
    else:
        print("Failed to create document.")