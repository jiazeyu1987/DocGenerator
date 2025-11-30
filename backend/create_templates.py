#!/usr/bin/env python3
"""
创建DOCX模板文件
需要安装：pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def create_basic_template():
    """创建基础文档模板"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # 标题1样式
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Microsoft YaHei'
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    # 标题2样式
    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Microsoft YaHei'
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    # 添加示例内容
    doc.add_heading('文档标题', 1)
    doc.add_paragraph('这是一个基础模板，适用于一般的文档转换。')
    doc.add_paragraph('支持中英文混合排版。')

    doc.add_heading('二级标题', 2)
    doc.add_paragraph('正文内容示例...')

    return doc

def create_technical_template():
    """创建技术文档模板"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # 设置字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10)

    # 标题样式
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Microsoft YaHei'
    heading1.font.size = Pt(18)
    heading1.font.bold = True

    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Microsoft YaHei'
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    # 添加示例
    doc.add_heading('技术文档模板', 1)
    doc.add_paragraph('适用于技术规范、API文档、代码注释等。')

    doc.add_heading('接口说明', 2)
    doc.add_paragraph('GET /api/endpoint')
    doc.add_paragraph('描述：获取用户信息')

    return doc

def create_business_template():
    """创建商务文档模板"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(2)
    section.right_margin = Inches(1.5)

    # 设置字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题样式
    heading1 = doc.styles['Heading 1']
    heading1.font.name = '黑体'
    heading1.font.size = Pt(22)
    heading1.font.bold = True
    heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading2 = doc.styles['Heading 2']
    heading2.font.name = '黑体'
    heading2.font.size = Pt(16)
    heading2.font.bold = True

    # 添加示例
    title = doc.add_heading('商务文档模板', 1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    doc.add_paragraph('适用于商业计划书、项目提案、合同文档等正式场合。')

    doc.add_heading('项目概述', 2)
    doc.add_paragraph('本部分介绍项目的基本情况...')

    return doc

def create_academic_template():
    """创建学术论文模板"""
    doc = Document()

    # 设置页面 (A4纸张)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)

    # 设置字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 标题样式
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Times New Roman'
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    # 添加示例
    title = doc.add_heading('学术论文模板', 1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    doc.add_paragraph('作者：张三')
    doc.add_paragraph('单位：某某大学')
    doc.add_paragraph('')

    doc.add_heading('摘要', 2)
    doc.add_paragraph('这是一个学术论文模板，适用于期刊论文、学位论文等学术写作。')

    doc.add_heading('关键词', 2)
    doc.add_paragraph('模板；学术论文；格式')

    return doc

def main():
    """创建所有模板"""
    templates_dir = "templates_store"
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)

    # 创建模板列表
    templates = [
        ("基础模板.docx", create_basic_template),
        ("技术文档模板.docx", create_technical_template),
        ("商务文档模板.docx", create_business_template),
        ("学术论文模板.docx", create_academic_template),
    ]

    print("正在创建DOCX模板...")

    try:
        for filename, create_func in templates:
            doc = create_func()
            filepath = os.path.join(templates_dir, filename)
            doc.save(filepath)
            print(f"✓ 已创建: {filename}")

        print(f"\n🎉 模板创建完成！")
        print(f"模板位置: {os.path.abspath(templates_dir)}")
        print("重启后端服务后即可在前端使用这些模板。")

    except Exception as e:
        print(f"❌ 创建模板时出错: {e}")
        print("请确保已安装 python-docx: pip install python-docx")

if __name__ == "__main__":
    main()